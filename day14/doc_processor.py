"""
doc_processor.py
================
Handles ingestion of PDF, DOCX, HTML, and Markdown files into
LangChain Document objects with rich metadata extraction.
"""

import os
import re
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any

# LangChain's base Document object — think of it as a dataclass
# with two fields: page_content (str) and metadata (dict)
# NEW (correct)
from langchain_core.documents import Document

# --- PDF Libraries ---
import pdfplumber          # Great for tables and structured PDFs
import fitz                # PyMuPDF — faster, better for image-heavy PDFs

# --- DOCX Library ---
from docx import Document as DocxDocument

# --- HTML Library ---
from bs4 import BeautifulSoup


# =============================================================================
# UTILITY: Metadata Builder
# =============================================================================

def build_base_metadata(file_path: str, loader_used: str) -> Dict[str, Any]:
    """
    Extracts filesystem-level metadata common to all file types.
    This runs for every file before type-specific processing begins.
    """
    path = Path(file_path)
    stat = path.stat()  # OS-level file stats

    return {
        "source": str(path.resolve()),          # Absolute path — critical for tracing answers
        "file_name": path.name,                  # e.g., "report.pdf"
        "file_type": path.suffix.lower(),        # e.g., ".pdf"
        "file_size_kb": round(stat.st_size / 1024, 2),
        "loader_used": loader_used,
        "ingestion_timestamp": datetime.now().isoformat(),
    }


# =============================================================================
# PDF PROCESSOR (Dual-Strategy: pdfplumber + PyMuPDF)
# =============================================================================

def load_pdf_pdfplumber(file_path: str) -> List[Document]:
    """
    Strategy 1: pdfplumber
    ----------------------
    Best for: Text-heavy PDFs, PDFs with tables, structured layouts.
    pdfplumber gives you fine-grained control over character-level extraction
    and can extract tables as structured data.
    
    How it works: Opens the PDF, iterates page-by-page, extracts text and
    any table structures found on that page.
    """
    documents = []
    base_meta = build_base_metadata(file_path, "pdfplumber")

    with pdfplumber.open(file_path) as pdf:
        
        # Extract document-level metadata from the PDF's own header
        pdf_meta = pdf.metadata or {}
        doc_level_meta = {
            "pdf_author": pdf_meta.get("Author", "Unknown"),
            "pdf_title": pdf_meta.get("Title", Path(file_path).stem),
            "pdf_creator": pdf_meta.get("Creator", "Unknown"),
            "total_pages": len(pdf.pages),
        }

        for page_num, page in enumerate(pdf.pages):
            # Extract raw text from this page
            text = page.extract_text()
            
            # Extract any tables on this page and append as text
            tables = page.extract_tables()
            table_text = ""
            if tables:
                for table in tables:
                    # Convert table rows to pipe-separated strings (markdown-ish format)
                    for row in table:
                        # Filter None values from cells
                        clean_row = [str(cell) if cell else "" for cell in row]
                        table_text += " | ".join(clean_row) + "\n"

            # Combine prose text with table text
            full_text = (text or "") + ("\n\nTABLES:\n" + table_text if table_text else "")

            if full_text.strip():  # Skip blank pages
                page_meta = {
                    **base_meta,
                    **doc_level_meta,
                    "page_number": page_num + 1,      # Human-readable (1-indexed)
                    "has_tables": bool(tables),
                    "char_count": len(full_text),
                }
                documents.append(Document(page_content=full_text.strip(), metadata=page_meta))

    print(f"[pdfplumber] Loaded {len(documents)} pages from {Path(file_path).name}")
    return documents


def load_pdf_pymupdf(file_path: str) -> List[Document]:
    """
    Strategy 2: PyMuPDF (fitz)
    --------------------------
    Best for: Scanned PDFs, PDFs with images, faster bulk processing.
    PyMuPDF can also extract images, links, and annotations — useful for
    multi-modal RAG applications.
    """
    documents = []
    base_meta = build_base_metadata(file_path, "pymupdf")

    # fitz.open() returns a Document object (confusingly, not a LangChain Document)
    pdf = fitz.open(file_path)

    # PyMuPDF's metadata is stored in pdf.metadata dict
    doc_level_meta = {
        "pdf_author": pdf.metadata.get("author", "Unknown"),
        "pdf_title": pdf.metadata.get("title", Path(file_path).stem),
        "total_pages": pdf.page_count,
        "is_encrypted": pdf.is_encrypted,
    }

    for page_num in range(pdf.page_count):
        page = pdf[page_num]  # Access page by index
        
        # get_text("text") returns plain text. Other options: "html", "dict", "json"
        text = page.get_text("text")
        
        # Extract hyperlinks from this page
        links = [link.get("uri", "") for link in page.get_links() if link.get("uri")]

        if text.strip():
            page_meta = {
                **base_meta,
                **doc_level_meta,
                "page_number": page_num + 1,
                "links_found": links,
                "link_count": len(links),
                "char_count": len(text),
            }
            documents.append(Document(page_content=text.strip(), metadata=page_meta))

    pdf.close()
    print(f"[PyMuPDF] Loaded {len(documents)} pages from {Path(file_path).name}")
    return documents


def load_pdf(file_path: str, strategy: str = "pdfplumber") -> List[Document]:
    """
    Router function: selects the PDF loading strategy.
    Defaults to pdfplumber; falls back to pymupdf on failure.
    """
    try:
        if strategy == "pdfplumber":
            return load_pdf_pdfplumber(file_path)
        else:
            return load_pdf_pymupdf(file_path)
    except Exception as e:
        print(f"[PDF Loader] Primary strategy failed: {e}. Trying PyMuPDF fallback...")
        return load_pdf_pymupdf(file_path)


# =============================================================================
# DOCX PROCESSOR
# =============================================================================

def load_docx(file_path: str) -> List[Document]:
    """
    DOCX Processor
    --------------
    python-docx reads .docx files as a tree of paragraphs and tables.
    We group paragraphs by heading structure to preserve document hierarchy —
    this is better than treating each paragraph as an independent chunk.
    
    Architecture decision: We return ONE Document per section (defined by
    heading boundaries), not one Document per paragraph. This preserves 
    semantic context.
    """
    base_meta = build_base_metadata(file_path, "python-docx")
    doc = DocxDocument(file_path)

    # Extract core document properties (author, dates, etc.)
    core_props = doc.core_properties
    doc_level_meta = {
        "doc_author": core_props.author or "Unknown",
        "doc_title": core_props.title or Path(file_path).stem,
        "doc_created": str(core_props.created) if core_props.created else "Unknown",
        "doc_modified": str(core_props.modified) if core_props.modified else "Unknown",
        "doc_subject": core_props.subject or "",
        "doc_keywords": core_props.keywords or "",
    }

    documents = []
    current_section_text = []
    current_heading = "Introduction"
    section_index = 0

    for para in doc.paragraphs:
        # paragraph.style.name is 'Heading 1', 'Heading 2', 'Normal', etc.
        if para.style.name.startswith("Heading"):
            # Save the previous section before starting a new one
            if current_section_text:
                full_text = "\n".join(current_section_text).strip()
                if full_text:
                    documents.append(Document(
                        page_content=full_text,
                        metadata={
                            **base_meta,
                            **doc_level_meta,
                            "section_title": current_heading,
                            "section_index": section_index,
                            "heading_level": para.style.name,
                        }
                    ))
                    section_index += 1
            # Start new section
            current_heading = para.text.strip()
            current_section_text = [current_heading]
        else:
            if para.text.strip():  # Skip empty paragraphs
                current_section_text.append(para.text.strip())

    # Don't forget the last section
    if current_section_text:
        full_text = "\n".join(current_section_text).strip()
        if full_text:
            documents.append(Document(
                page_content=full_text,
                metadata={**base_meta, **doc_level_meta,
                           "section_title": current_heading,
                           "section_index": section_index}
            ))

    # Also extract table content
    for table_idx, table in enumerate(doc.tables):
        rows_text = []
        for row in table.rows:
            cell_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cell_texts:
                rows_text.append(" | ".join(cell_texts))
        if rows_text:
            table_content = "\n".join(rows_text)
            documents.append(Document(
                page_content=table_content,
                metadata={
                    **base_meta,
                    **doc_level_meta,
                    "section_title": f"Table {table_idx + 1}",
                    "is_table": True,
                    "table_index": table_idx,
                }
            ))

    print(f"[DOCX] Loaded {len(documents)} sections from {Path(file_path).name}")
    return documents


# =============================================================================
# HTML PROCESSOR
# =============================================================================

def load_html(file_path: str) -> List[Document]:
    """
    HTML Processor
    --------------
    BeautifulSoup parses the HTML DOM. We extract:
    1. Metadata from <meta> tags and <title>
    2. Content by walking heading tags (h1-h6) to preserve structure
    3. We strip scripts, styles, and nav elements — they're noise for RAG.
    """
    base_meta = build_base_metadata(file_path, "beautifulsoup4")

    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        html_content = f.read()

    # 'html.parser' is Python's built-in parser — no extra install needed
    soup = BeautifulSoup(html_content, "html.parser")

    # --- Extract HTML-level metadata ---
    html_meta = {}
    for tag in soup.find_all("meta"):
        name = tag.get("name", tag.get("property", ""))
        content = tag.get("content", "")
        if name and content:
            html_meta[f"html_meta_{name}"] = content

    title_tag = soup.find("title")
    html_title = title_tag.get_text(strip=True) if title_tag else Path(file_path).stem
    html_meta["html_title"] = html_title

    # --- Remove noisy tags BEFORE extraction ---
    for noisy in soup(["script", "style", "nav", "footer", "header", "aside"]):
        noisy.decompose()  # Removes element from the tree entirely

    documents = []
    headings = soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"])

    if headings:
        # Section-based extraction: collect all siblings until the next heading
        for idx, heading in enumerate(headings):
            section_text = [heading.get_text(strip=True)]
            
            # Walk siblings until we hit the next heading
            for sibling in heading.find_next_siblings():
                if sibling.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                    break  # Stop at the next heading
                text = sibling.get_text(separator=" ", strip=True)
                if text:
                    section_text.append(text)

            full_text = "\n".join(section_text).strip()
            if full_text:
                documents.append(Document(
                    page_content=full_text,
                    metadata={
                        **base_meta,
                        **html_meta,
                        "section_heading": heading.get_text(strip=True),
                        "heading_tag": heading.name,  # 'h1', 'h2', etc.
                        "section_index": idx,
                    }
                ))
    else:
        # Fallback: no headings found, extract all body text
        body = soup.find("body")
        text = body.get_text(separator="\n", strip=True) if body else soup.get_text()
        if text:
            documents.append(Document(
                page_content=text,
                metadata={**base_meta, **html_meta, "section_heading": "body"}
            ))

    print(f"[HTML] Loaded {len(documents)} sections from {Path(file_path).name}")
    return documents


# =============================================================================
# MARKDOWN PROCESSOR
# =============================================================================

def load_markdown(file_path: str) -> List[Document]:
    """
    Markdown Processor
    ------------------
    Markdown is the cleanest format for RAG because structure is explicit.
    We split on heading lines (lines starting with #) to create sections,
    and extract frontmatter metadata (YAML between --- delimiters).
    """
    base_meta = build_base_metadata(file_path, "custom-markdown")

    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()

    # --- Extract YAML frontmatter ---
    frontmatter_meta = {}
    frontmatter_pattern = re.compile(r"^---\s*\n(.*?)\n---\s*\n", re.DOTALL)
    fm_match = frontmatter_pattern.match(content)
    if fm_match:
        fm_text = fm_match.group(1)
        content = content[fm_match.end():]  # Remove frontmatter from content
        # Parse simple key: value pairs (not full YAML for simplicity)
        for line in fm_text.strip().split("\n"):
            if ":" in line:
                key, _, value = line.partition(":")
                frontmatter_meta[f"fm_{key.strip()}"] = value.strip()

    # --- Split by heading lines ---
    # re.split with a capture group keeps the delimiter in the result list
    sections = re.split(r"(\n#{1,6} .+)", content)

    documents = []
    current_heading = "Introduction"
    current_content = []
    section_index = 0

    for part in sections:
        heading_match = re.match(r"\n(#{1,6}) (.+)", part)
        if heading_match:
            # Save previous section
            if current_content:
                text = "\n".join(current_content).strip()
                if text:
                    level = current_heading.count("#", 0, 7) if current_heading.startswith("#") else 0
                    documents.append(Document(
                        page_content=text,
                        metadata={
                            **base_meta,
                            **frontmatter_meta,
                            "section_heading": current_heading,
                            "heading_level": level,
                            "section_index": section_index,
                        }
                    ))
                    section_index += 1
            # Update current heading
            hashes = heading_match.group(1)
            heading_text = heading_match.group(2).strip()
            current_heading = heading_text
            current_content = [f"{hashes} {heading_text}"]
        else:
            if part.strip():
                current_content.append(part)

    # Last section
    if current_content:
        text = "\n".join(current_content).strip()
        if text:
            documents.append(Document(
                page_content=text,
                metadata={
                    **base_meta, **frontmatter_meta,
                    "section_heading": current_heading,
                    "section_index": section_index,
                }
            ))

    print(f"[Markdown] Loaded {len(documents)} sections from {Path(file_path).name}")
    return documents


# =============================================================================
# MASTER DISPATCHER
# =============================================================================

def load_document(file_path: str, pdf_strategy: str = "pdfplumber") -> List[Document]:
    """
    Single entry point: inspects the file extension and routes to the
    correct loader. Returns a list of LangChain Document objects.
    
    Usage:
        docs = load_document("report.pdf")
        docs = load_document("notes.md")
    """
    ext = Path(file_path).suffix.lower()
    loaders = {
        ".pdf":  lambda: load_pdf(file_path, strategy=pdf_strategy),
        ".docx": lambda: load_docx(file_path),
        ".html": lambda: load_html(file_path),
        ".htm":  lambda: load_html(file_path),
        ".md":   lambda: load_markdown(file_path),
        ".txt":  lambda: [Document(
                    page_content=open(file_path, "r").read(),
                    metadata=build_base_metadata(file_path, "plain-text")
                 )],
    }
    
    if ext not in loaders:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {list(loaders.keys())}")
    
    return loaders[ext]()


def load_directory(dir_path: str, pdf_strategy: str = "pdfplumber") -> List[Document]:
    """
    Recursively loads all supported documents from a directory.
    Useful for bulk ingestion.
    """
    all_docs = []
    supported_extensions = {".pdf", ".docx", ".html", ".htm", ".md", ".txt"}
    
    for path in Path(dir_path).rglob("*"):
        if path.suffix.lower() in supported_extensions:
            try:
                docs = load_document(str(path), pdf_strategy=pdf_strategy)
                all_docs.extend(docs)
            except Exception as e:
                print(f"[WARNING] Failed to load {path.name}: {e}")
    
    print(f"\n[DocProcessor] Total documents loaded: {len(all_docs)}")
    return all_docs


# =============================================================================
# STANDALONE TEST
# =============================================================================

if __name__ == "__main__":
    import json

    # Create sample files for testing
    os.makedirs("sample_docs", exist_ok=True)

    # Create a sample markdown file to test with
    sample_md = """---
title: RAG Tutorial
author: AlgoProfessor
date: 2026-01-01
---

# Introduction to RAG

Retrieval-Augmented Generation (RAG) combines retrieval systems with LLMs.

## How it Works

The system retrieves relevant documents and uses them as context.

### Step 1: Ingestion

Documents are loaded, chunked, and embedded into a vector store.

### Step 2: Retrieval

User queries are embedded and matched against stored vectors.

## Benefits of RAG

RAG reduces hallucination and allows LLMs to access current information.
"""
    with open("sample_docs/sample.md", "w") as f:
        f.write(sample_md)

    # Create a sample HTML file
    sample_html = """<!DOCTYPE html>
<html>
<head>
    <title>RAG Systems Overview</title>
    <meta name="author" content="AlgoProfessor">
    <meta name="description" content="Overview of RAG systems">
</head>
<body>
    <nav>Navigation here</nav>
    <h1>RAG Systems</h1>
    <p>RAG stands for Retrieval-Augmented Generation.</p>
    <h2>Architecture</h2>
    <p>The architecture consists of three main components: ingestion, retrieval, and generation.</p>
    <h2>Evaluation</h2>
    <p>RAGAS is the standard framework for evaluating RAG pipelines.</p>
</body>
</html>"""
    with open("sample_docs/sample.html", "w") as f:
        f.write(sample_html)

    # Test loading
    print("=" * 60)
    print("Testing doc_processor.py")
    print("=" * 60)

    md_docs = load_document("sample_docs/sample.md")
    print(f"\nMarkdown loaded {len(md_docs)} documents")
    for doc in md_docs:
        print(f"  Section: '{doc.metadata['section_heading']}' | Chars: {len(doc.page_content)}")

    html_docs = load_document("sample_docs/sample.html")
    print(f"\nHTML loaded {len(html_docs)} documents")
    for doc in html_docs:
        print(f"  Section: '{doc.metadata['section_heading']}' | Tag: {doc.metadata['heading_tag']}")

    print("\n[SUCCESS] doc_processor.py is working correctly!")